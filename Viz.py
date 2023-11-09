import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def plot_covid_pie_chart(datafile, output_filename):
    # Load data from CSV file and remove commas from numeric columns
    df = pd.read_csv(datafile, thousands=',')

    # Get data for pie chart
    countries = df['Country']
    total_cases = df['Total Cases']

    plt.figure(figsize=(8, 8))
    plt.pie(total_cases, labels=countries, autopct='%1.1f%%', shadow=True, startangle=140)
    plt.title('COVID-19 Total Cases Distribution in Europe')
    plt.savefig('covid_pie_chart.png')
    plt.close()

    doc = Document()
    doc.add_heading('Data Visualizations Report', level=1)
    doc.add_heading('COVID-19 Total Cases Distribution', level=2)
    doc.add_picture('covid_pie_chart.png', width=Inches(6))
    doc.add_paragraph('Figure 1: Distribution of COVID-19 Total Cases in Europe.')
    doc.save(output_filename)

def plot_stock_prices_line_chart(datafiles, output_filename):
    merged_data = {}  # Dictionary to store merged data for each company
    for company, file_path in datafiles.items():
        df = pd.read_csv(file_path)
        merged_data[company] = df[['Date', 'Close']].rename(columns={'Close': f'{company}'})

    # Merge data from all companies
    merged_data_all = merged_data[list(merged_data.keys())[0]]
    for company in merged_data.keys():
        merged_data_all = pd.merge(merged_data_all, merged_data[company], on='Date', how='inner')

    plt.figure(figsize=(10, 6))
    for company in merged_data_all.columns[1:]:
        plt.plot(merged_data_all['Date'], merged_data_all[company], marker='o', label=company)

    plt.xlabel('Date')
    plt.ylabel('Close Price')
    plt.title('Unit share Prices of Oil Companies between before and after ukrain war')
    plt.xticks(rotation=45)
    plt.legend()
    plt.tight_layout()
    plt.savefig('merged_stock_prices_line_chart.png')
    plt.close()

    doc = Document(output_filename)
    doc.add_heading('Prices of Oil Companies', level=2)
    doc.add_picture('merged_stock_prices_line_chart.png', width=Inches(6))
    doc.add_paragraph('Figure: Merged Close Prices of Oil Companies.')
    doc.save(output_filename)

def plot_boundary_runs_vs_non_boundary_runs(datafile, output_filename, top_n=10):
    # Load data from CSV file
    mens_odi_stats = pd.read_csv(datafile)

    # Convert relevant columns to numeric, coerce errors to NaN
    mens_odi_stats['Innings Runs Scored Num'] = pd.to_numeric(mens_odi_stats['Innings Runs Scored Num'], errors='coerce')
    mens_odi_stats['Innings Boundary Fours'] = pd.to_numeric(mens_odi_stats['Innings Boundary Fours'], errors='coerce')
    mens_odi_stats['Innings Boundary Sixes'] = pd.to_numeric(mens_odi_stats['Innings Boundary Sixes'], errors='coerce')

    # Drop rows with NaN values in relevant columns
    mens_odi_stats.dropna(subset=['Innings Runs Scored Num', 'Innings Boundary Fours', 'Innings Boundary Sixes'], inplace=True)

    # Calculate boundary runs and non-boundary runs
    boundary_runs = mens_odi_stats['Innings Boundary Fours'] * 4 + mens_odi_stats['Innings Boundary Sixes'] * 6
    non_boundary_runs = mens_odi_stats['Innings Runs Scored Num'] - boundary_runs

    # Add boundary_runs and non_boundary_runs as new columns
    mens_odi_stats['Boundary Runs'] = boundary_runs
    mens_odi_stats['Non-Boundary Runs'] = non_boundary_runs

    # Group data by player and calculate total boundary runs and non-boundary runs
    total_boundary_runs = mens_odi_stats.groupby('Innings Player')['Boundary Runs'].sum()
    total_non_boundary_runs = mens_odi_stats.groupby('Innings Player')['Non-Boundary Runs'].sum()

    # Get top n players by total runs scored
    top_n_players = total_boundary_runs.add(total_non_boundary_runs).nlargest(top_n)

    # Plot stacked bar chart for top n players by boundary and non-boundary runs
    plt.figure(figsize=(10, 6))
    plt.bar(top_n_players.index, total_boundary_runs[top_n_players.index], label='Boundary Runs', color='skyblue')
    plt.bar(top_n_players.index, total_non_boundary_runs[top_n_players.index], label='Non-Boundary Runs', color='orange', bottom=total_boundary_runs[top_n_players.index])
    plt.xlabel('Player')
    plt.ylabel('Runs Scored')
    plt.title(f'Boundary Runs vs Non-Boundary Runs for Top {top_n} ODI Players')
    plt.xticks(rotation=45)
    plt.legend()
    plt.tight_layout()
    plt.savefig('boundary_runs_vs_non_boundary_runs.png')
    plt.close()

    doc = Document(output_filename)
    doc.add_heading(f'Boundary Runs vs Non-Boundary Runs for Top {top_n} ODI Players', level=2)
    doc.add_picture('boundary_runs_vs_non_boundary_runs.png', width=Inches(6))
    doc.add_paragraph(f'Figure: Boundary Runs vs Non-Boundary Runs for Top {top_n} ODI Players.')
    doc.save(output_filename)

if __name__ == "__main__":
    covid_datafile = "Covid.csv"
    stock_prices_datafiles = {
        'Chevron': 'CVX.csv',
        'Shell': 'SHEL.csv',
        'ExxonMobil': 'XOM.csv'
    }
    boundary_runs_datafile = "mens_odi_stats.csv"
    output_filename = "DataVisualizationsReport.docx"

    plot_covid_pie_chart(covid_datafile, output_filename)
    plot_stock_prices_line_chart(stock_prices_datafiles, output_filename)
    plot_boundary_runs_vs_non_boundary_runs(boundary_runs_datafile, output_filename, top_n=10)